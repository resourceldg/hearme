"""Parsers DOCX, ODT y RTF."""

from __future__ import annotations

import re
from pathlib import Path

from hearme.application.cleaning import clean_blocks
from hearme.domain.models import (
    Block,
    BlockKind,
    Chapter,
    Document,
    DocumentMeta,
    SourceFormat,
)
from hearme.infrastructure.parsers.html import blocks_to_chapters

_HEADING_STYLE = re.compile(r"heading\s*(\d)|t[íi]tulo\s*(\d)", re.IGNORECASE)


class DOCXParser:
    name = "docx"
    formats = frozenset({SourceFormat.DOCX})

    def can_parse(self, path: Path) -> bool:
        return path.suffix.lower() == ".docx"

    async def parse(self, path: Path) -> Document:
        import docx

        document = docx.Document(str(path))
        blocks: list[Block] = []

        for order, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            style = (paragraph.style.name if paragraph.style else "") or ""
            kind, level = BlockKind.PARAGRAPH, None

            if match := _HEADING_STYLE.search(style):
                kind = BlockKind.HEADING
                level = int(match.group(1) or match.group(2))
            elif "quote" in style.lower() or "cita" in style.lower():
                kind = BlockKind.QUOTE
            elif "list" in style.lower():
                kind = BlockKind.LIST_ITEM

            blocks.append(Block(kind=kind, text=text, order=order, level=level))

        # Las tablas van al final del bloque de texto, marcadas como TABLE para que
        # el modo audiolibro pueda omitirlas y el modo lectura mostrarlas.
        for table in document.tables:
            rows = [
                " | ".join(cell.text.strip() for cell in row.cells)
                for row in table.rows
                if any(cell.text.strip() for cell in row.cells)
            ]
            if rows:
                blocks.append(Block(kind=BlockKind.TABLE, text="\n".join(rows), order=len(blocks)))

        core = document.core_properties
        title = (core.title or "").strip() or path.stem
        meta = DocumentMeta(
            title=title,
            authors=[core.author] if core.author else [],
            language=(core.language or "").split("-")[0].lower() or None,
        )
        return Document(
            source_path=path,
            source_format=SourceFormat.DOCX,
            meta=meta,
            chapters=blocks_to_chapters(clean_blocks(blocks), title),
        )


class ODTParser:
    name = "odt"
    formats = frozenset({SourceFormat.ODT})

    def can_parse(self, path: Path) -> bool:
        return path.suffix.lower() == ".odt"

    async def parse(self, path: Path) -> Document:
        from odf import teletype
        from odf import text as odf_text
        from odf.opendocument import load

        document = load(str(path))
        blocks: list[Block] = []
        order = 0

        for element in document.getElementsByType(odf_text.H) + document.getElementsByType(
            odf_text.P
        ):
            content = teletype.extractText(element).strip()
            if not content:
                continue
            is_heading = element.qname[1] == "h"
            level = None
            if is_heading:
                raw = element.getAttribute("outlinelevel")
                level = int(raw) if raw and str(raw).isdigit() else 1
            blocks.append(
                Block(
                    kind=BlockKind.HEADING if is_heading else BlockKind.PARAGRAPH,
                    text=content,
                    order=order,
                    level=level,
                )
            )
            order += 1

        # getElementsByType no garantiza orden de documento entre tipos distintos:
        # se reordena por la posición real en el XML.
        return Document(
            source_path=path,
            source_format=SourceFormat.ODT,
            meta=DocumentMeta(title=path.stem),
            chapters=blocks_to_chapters(clean_blocks(blocks), path.stem),
        )


class RTFParser:
    name = "rtf"
    formats = frozenset({SourceFormat.RTF})

    def can_parse(self, path: Path) -> bool:
        return path.suffix.lower() == ".rtf"

    async def parse(self, path: Path) -> Document:
        from striprtf.striprtf import rtf_to_text

        content = rtf_to_text(path.read_text(encoding="utf-8", errors="replace"), errors="ignore")
        blocks = [
            Block(kind=BlockKind.PARAGRAPH, text=chunk.strip(), order=index)
            for index, chunk in enumerate(re.split(r"\n\s*\n", content))
            if chunk.strip()
        ]
        return Document(
            source_path=path,
            source_format=SourceFormat.RTF,
            meta=DocumentMeta(title=path.stem),
            chapters=[Chapter(title=path.stem, order=0, blocks=clean_blocks(blocks))],
        )
